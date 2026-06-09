import os
import shutil
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List
import aiofiles
from docx import Document
from web.file_logger import file_logger

class FileOperationManager:
    def __init__(self, upload_dir: str = "C:/Users/hp/Desktop/upload"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        print(f"[FileManager] 文件操作目录: {self.upload_dir.absolute()}")
    
    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名，确保有 .txt 扩展名"""
        # 去掉路径分隔符，只取最后一部分
        filename = filename.replace('\\', '/').split('/')[-1]
        # 去掉可能的盘符前缀
        if ':' in filename:
            filename = filename.split(':')[-1]
        # 去掉开头的特殊字符
        filename = filename.strip().strip('/').strip('\\')
        # 如果没有扩展名，默认加 .txt
        if '.' not in filename:
            filename = filename + '.txt'
        return filename
    
    def _get_file_path(self, filename: str) -> Path:
        """获取文件完整路径"""
        safe_name = self._sanitize_filename(filename)
        return self.upload_dir / safe_name
    
    def _fuzzy_match(self, filename: str) -> Optional[Path]:
        """模糊匹配文件名"""
        name_lower = filename.lower()
        for f in self.upload_dir.iterdir():
            if f.is_file():
                f_name_lower = f.name.lower()
                # 完全匹配
                if f_name_lower == name_lower:
                    return f
                # 包含匹配（文件名包含关键词）
                if name_lower in f_name_lower:
                    return f
        return None
    
    async def create_file(self, filename: str, content: str) -> Dict:
        """创建新文件"""
        start_time = datetime.now()
        try:
            safe_name = self._sanitize_filename(filename)
            file_path = self._get_file_path(safe_name)
            
            if file_path.exists():
                error_msg = f"文件 {safe_name} 已存在"
                file_logger.log_operation(
                    operation="create_file",
                    success=False,
                    details={
                        "filename": safe_name,
                        "content_length": len(content)
                    },
                    error=error_msg
                )
                return {"success": False, "message": error_msg}
            
            file_type = 'txt' if safe_name.endswith('.txt') else 'docx'
            
            if file_type == 'txt':
                async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
                    await f.write(content)
            else:
                doc = Document()
                doc.add_paragraph(content)
                doc.save(file_path)
            
            file_stats = file_path.stat()
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="create_file",
                success=True,
                details={
                    "filename": safe_name,
                    "file_type": file_type,
                    "content_length": len(content),
                    "size_bytes": file_stats.st_size,
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"文件 {safe_name} 创建成功",
                "data": {
                    "name": safe_name,
                    "type": "file",
                    "size": file_stats.st_size,
                    "created_at": datetime.fromtimestamp(file_stats.st_ctime),
                    "updated_at": datetime.fromtimestamp(file_stats.st_mtime)
                }
            }
        except Exception as e:
            error_msg = str(e)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="create_file",
                success=False,
                details={
                    "filename": filename,
                    "content_length": len(content) if content else 0,
                    "duration_ms": round(duration_ms, 2)
                },
                error=error_msg
            )
            return {"success": False, "message": f"创建文件失败: {error_msg}"}
    
    async def read_file(self, filename: str) -> Dict:
        """读取文件内容"""
        start_time = datetime.now()
        try:
            safe_name = self._sanitize_filename(filename)
            file_path = self._get_file_path(safe_name)
            
            if not file_path.exists():
                matched = self._fuzzy_match(safe_name)
                if matched:
                    file_path = matched
                    safe_name = file_path.name
                else:
                    error_msg = f"文件 {safe_name} 不存在"
                    file_logger.log_operation(
                        operation="read_file",
                        success=False,
                        details={"filename": safe_name},
                        error=error_msg
                    )
                    return {"success": False, "message": error_msg}
            
            file_type = 'txt' if safe_name.endswith('.txt') else 'docx'
            
            if file_type == 'txt':
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    content = await f.read()
            else:
                doc = Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            file_stats = file_path.stat()
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="read_file",
                success=True,
                details={
                    "filename": safe_name,
                    "file_type": file_type,
                    "content_length": len(content),
                    "size_bytes": file_stats.st_size,
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"读取文件 {safe_name} 成功",
                "data": {
                    "name": safe_name,
                    "type": "file",
                    "content": content,
                    "size": file_stats.st_size,
                    "updated_at": datetime.fromtimestamp(file_stats.st_mtime)
                }
            }
        except Exception as e:
            error_msg = str(e)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="read_file",
                success=False,
                details={"filename": filename, "duration_ms": round(duration_ms, 2)},
                error=error_msg
            )
            return {"success": False, "message": f"读取文件失败: {error_msg}"}
    
    async def update_file(self, filename: str, new_content: str) -> Dict:
        """更新文件内容"""
        start_time = datetime.now()
        try:
            safe_name = self._sanitize_filename(filename)
            file_path = self._get_file_path(safe_name)
            
            if not file_path.exists():
                matched = self._fuzzy_match(safe_name)
                if matched:
                    file_path = matched
                    safe_name = file_path.name
                else:
                    error_msg = f"文件 {safe_name} 不存在"
                    file_logger.log_operation(
                        operation="update_file",
                        success=False,
                        details={"filename": safe_name},
                        error=error_msg
                    )
                    return {"success": False, "message": error_msg}
            
            old_size = file_path.stat().st_size
            
            file_type = 'txt' if safe_name.endswith('.txt') else 'docx'
            
            if file_type == 'txt':
                async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
                    await f.write(new_content)
            else:
                doc = Document()
                doc.add_paragraph(new_content)
                doc.save(file_path)
            
            file_stats = file_path.stat()
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="update_file",
                success=True,
                details={
                    "filename": safe_name,
                    "file_type": file_type,
                    "old_size_bytes": old_size,
                    "new_size_bytes": file_stats.st_size,
                    "content_length": len(new_content),
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"文件 {safe_name} 更新成功",
                "data": {
                    "name": safe_name,
                    "type": "file",
                    "size": file_stats.st_size,
                    "updated_at": datetime.fromtimestamp(file_stats.st_mtime)
                }
            }
        except Exception as e:
            error_msg = str(e)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="update_file",
                success=False,
                details={
                    "filename": filename,
                    "content_length": len(new_content) if new_content else 0,
                    "duration_ms": round(duration_ms, 2)
                },
                error=error_msg
            )
            return {"success": False, "message": f"更新文件失败: {error_msg}"}
    
    async def rename_file(self, old_name: str, new_name: str) -> Dict:
        """重命名文件"""
        start_time = datetime.now()
        try:
            safe_old = self._sanitize_filename(old_name)
            old_path = self._get_file_path(safe_old)
            
            if not old_path.exists():
                matched = self._fuzzy_match(safe_old)
                if matched:
                    old_path = matched
                    safe_old = old_path.name
                else:
                    error_msg = f"文件 {safe_old} 不存在"
                    file_logger.log_operation(
                        operation="rename_file",
                        success=False,
                        details={"old_name": safe_old},
                        error=error_msg
                    )
                    return {"success": False, "message": error_msg}
            
            safe_new = self._sanitize_filename(new_name)
            new_path = self._get_file_path(safe_new)
            
            if new_path.exists():
                error_msg = f"文件 {safe_new} 已存在"
                file_logger.log_operation(
                    operation="rename_file",
                    success=False,
                    details={"old_name": safe_old, "new_name": safe_new},
                    error=error_msg
                )
                return {"success": False, "message": error_msg}
            
            old_path.rename(new_path)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="rename_file",
                success=True,
                details={
                    "old_name": safe_old,
                    "new_name": safe_new,
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"文件 {safe_old} 已重命名为 {safe_new}",
                "data": {
                    "old_name": safe_old,
                    "new_name": safe_new,
                    "type": "file"
                }
            }
        except Exception as e:
            error_msg = str(e)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="rename_file",
                success=False,
                details={
                    "old_name": old_name,
                    "new_name": new_name,
                    "duration_ms": round(duration_ms, 2)
                },
                error=error_msg
            )
            return {"success": False, "message": f"重命名失败: {error_msg}"}
    
    async def copy_file(self, source_name: str, dest_name: str) -> Dict:
        """复制文件（目标名自动规范化）"""
        start_time = datetime.now()
        try:
            safe_source = self._sanitize_filename(source_name)
            source_path = self._get_file_path(safe_source)
            
            if not source_path.exists():
                matched = self._fuzzy_match(safe_source)
                if matched:
                    source_path = matched
                    safe_source = source_path.name
                else:
                    error_msg = f"源文件 {safe_source} 不存在"
                    file_logger.log_operation(
                        operation="copy_file",
                        success=False,
                        details={"source": safe_source},
                        error=error_msg
                    )
                    return {"success": False, "message": error_msg}
            
            # 目标名也规范化（自动加 .txt）
            safe_dest = self._sanitize_filename(dest_name)
            dest_path = self._get_file_path(safe_dest)
            
            if dest_path.exists():
                error_msg = f"目标文件 {safe_dest} 已存在"
                file_logger.log_operation(
                    operation="copy_file",
                    success=False,
                    details={"source": safe_source, "destination": safe_dest},
                    error=error_msg
                )
                return {"success": False, "message": error_msg}
            
            shutil.copy2(source_path, dest_path)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="copy_file",
                success=True,
                details={
                    "source": safe_source,
                    "destination": safe_dest,
                    "source_size_bytes": source_path.stat().st_size,
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"文件 {safe_source} 已复制到 {safe_dest}",
                "data": {
                    "source": safe_source,
                    "destination": safe_dest,
                    "type": "file"
                }
            }
        except Exception as e:
            error_msg = str(e)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="copy_file",
                success=False,
                details={
                    "source": source_name,
                    "destination": dest_name,
                    "duration_ms": round(duration_ms, 2)
                },
                error=error_msg
            )
            return {"success": False, "message": f"复制文件失败: {error_msg}"}
    
    async def list_files(self) -> Dict:
        """列出所有文件"""
        start_time = datetime.now()
        try:
            files = []
            for file_path in self.upload_dir.iterdir():
                if file_path.is_file():
                    file_stats = file_path.stat()
                    files.append({
                        "name": file_path.name,
                        "type": "file",
                        "size": file_stats.st_size,
                        "created_at": datetime.fromtimestamp(file_stats.st_ctime),
                        "updated_at": datetime.fromtimestamp(file_stats.st_mtime)
                    })
            
            files.sort(key=lambda x: x.get("name", ""))
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="list_files",
                success=True,
                details={
                    "file_count": len(files),
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"找到 {len(files)} 个文件",
                "data": files
            }
        except Exception as e:
            error_msg = str(e)
            file_logger.log_operation(
                operation="list_files",
                success=False,
                details={},
                error=error_msg
            )
            return {"success": False, "message": f"列出文件失败: {error_msg}"}
    
    async def search_files(self, keyword: str) -> Dict:
        """搜索文件内容"""
        start_time = datetime.now()
        try:
            results = []
            for file_path in self.upload_dir.iterdir():
                if file_path.is_file():
                    try:
                        file_type = 'txt' if file_path.name.endswith('.txt') else 'docx'
                        
                        if file_type == 'txt':
                            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                                content = await f.read()
                        else:
                            doc = Document(file_path)
                            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                        
                        if keyword.lower() in content.lower():
                            idx = content.lower().find(keyword.lower())
                            start = max(0, idx - 50)
                            end = min(len(content), idx + 100)
                            preview = content[start:end]
                            if start > 0:
                                preview = "..." + preview
                            if end < len(content):
                                preview = preview + "..."
                            
                            results.append({
                                "name": file_path.name,
                                "type": "file",
                                "content_preview": preview
                            })
                    except Exception:
                        continue
            
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="search_files",
                success=True,
                details={
                    "keyword": keyword,
                    "match_count": len(results),
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"找到 {len(results)} 个包含关键词的文件",
                "data": results
            }
        except Exception as e:
            error_msg = str(e)
            file_logger.log_operation(
                operation="search_files",
                success=False,
                details={"keyword": keyword},
                error=error_msg
            )
            return {"success": False, "message": f"搜索文件失败: {error_msg}"}
    
    async def merge_files(self, source_names: List[str], dest_name: str, separator: str = "\n\n") -> Dict:
        """合并多个文件的内容到一个新文件"""
        start_time = datetime.now()
        try:
            contents = []
            file_list = []
            
            for source_name in source_names:
                safe_source = self._sanitize_filename(source_name)
                source_path = self._get_file_path(safe_source)
                
                if not source_path.exists():
                    matched = self._fuzzy_match(safe_source)
                    if matched:
                        source_path = matched
                        safe_source = source_path.name
                    else:
                        error_msg = f"源文件 {safe_source} 不存在"
                        file_logger.log_operation(
                            operation="merge_files",
                            success=False,
                            details={"sources": source_names, "missing": safe_source},
                            error=error_msg
                        )
                        return {"success": False, "message": error_msg}
                
                file_list.append(safe_source)
                
                file_type = 'txt' if safe_source.endswith('.txt') else 'docx'
                
                if file_type == 'txt':
                    async with aiofiles.open(source_path, 'r', encoding='utf-8') as f:
                        content = await f.read()
                else:
                    doc = Document(source_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                # 添加文件标题标记
                if len(source_names) > 1:
                    header = f"=== {safe_source} ===\n"
                    contents.append(header + content)
                else:
                    contents.append(content)
            
            # 合并内容
            merged_content = separator.join(contents)
            
            # 创建新文件
            safe_dest = self._sanitize_filename(dest_name)
            dest_path = self._get_file_path(safe_dest)
            
            if dest_path.exists():
                error_msg = f"目标文件 {safe_dest} 已存在"
                file_logger.log_operation(
                    operation="merge_files",
                    success=False,
                    details={"sources": file_list, "destination": safe_dest},
                    error=error_msg
                )
                return {"success": False, "message": error_msg}
            
            async with aiofiles.open(dest_path, 'w', encoding='utf-8') as f:
                await f.write(merged_content)
            
            file_stats = dest_path.stat()
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            file_logger.log_operation(
                operation="merge_files",
                success=True,
                details={
                    "sources": file_list,
                    "destination": safe_dest,
                    "source_count": len(file_list),
                    "total_size_bytes": file_stats.st_size,
                    "duration_ms": round(duration_ms, 2)
                }
            )
            
            return {
                "success": True,
                "message": f"已将 {', '.join(file_list)} 合并到 {safe_dest}",
                "data": {
                    "sources": file_list,
                    "destination": safe_dest,
                    "size": file_stats.st_size,
                    "created_at": datetime.fromtimestamp(file_stats.st_ctime)
                }
            }
            
        except Exception as e:
            error_msg = str(e)
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="merge_files",
                success=False,
                details={
                    "sources": source_names,
                    "destination": dest_name,
                    "duration_ms": round(duration_ms, 2)
                },
                error=error_msg
            )
            return {"success": False, "message": f"合并文件失败: {error_msg}"}
        
    async def clean_file(self, filepath: str, rules: Dict, command_text: str = None) -> Dict:
        """清洗文件（基础清洗）"""
        start_time = datetime.now()
        from tools.file_cleaner import file_cleaner
        
        try:
            result = file_cleaner.clean_file(filepath=filepath, rules=rules)
            
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            if result.get("success"):
                file_logger.log_operation(
                    operation="clean_file",
                    success=True,
                    details={
                        "filepath": result.get("filepath", filepath),
                        "original_length": result.get("original_length", 0),
                        "new_length": result.get("new_length", 0),
                        "original_lines": result.get("original_lines", 0),
                        "new_lines": result.get("new_lines", 0),
                        "rules": rules,
                        "duration_ms": round(duration_ms, 2)
                    },
                    command_text=command_text
                )
            else:
                file_logger.log_operation(
                    operation="clean_file",
                    success=False,
                    details={"filepath": filepath, "rules": rules, "duration_ms": round(duration_ms, 2)},
                    error=result.get("message", "清洗失败"),
                    command_text=command_text
                )
            
            return result
        except Exception as e:
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="clean_file",
                success=False,
                details={"filepath": filepath, "duration_ms": round(duration_ms, 2)},
                error=str(e),
                command_text=command_text
            )
            return {"success": False, "message": f"清洗失败: {str(e)}"}

    async def ai_clean_file(self, filepath: str, instruction: str, command_text: str = None) -> Dict:
        """AI 智能清洗文件"""
        start_time = datetime.now()
        from tools.file_cleaner import file_cleaner
        
        try:
            result = await file_cleaner.ai_clean_file(filepath=filepath, instruction=instruction)
            
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            
            if result.get("success"):
                file_logger.log_operation(
                    operation="ai_clean_file",
                    success=True,
                    details={
                        "filepath": result.get("filepath", filepath),
                        "original_length": result.get("original_length", 0),
                        "new_length": result.get("new_length", 0),
                        "instruction": instruction[:200] + "..." if len(instruction) > 200 else instruction,
                        "duration_ms": round(duration_ms, 2)
                    },
                    command_text=command_text
                )
                
                if result.get("target_filename"):
                    file_logger.log_operation(
                        operation="ai_clean_file_with_rename",
                        success=True,
                        details={
                            "source_file": filepath,
                            "target_file": result.get("filepath", ""),
                            "original_length": result.get("original_length", 0),
                            "new_length": result.get("new_length", 0),
                            "duration_ms": round(duration_ms, 2)
                        },
                        command_text=command_text
                    )
            else:
                file_logger.log_operation(
                    operation="ai_clean_file",
                    success=False,
                    details={"filepath": filepath, "duration_ms": round(duration_ms, 2)},
                    error=result.get("message", "AI清洗失败"),
                    command_text=command_text
                )
            
            return result
        except Exception as e:
            duration_ms = (datetime.now() - start_time).total_seconds() * 1000
            file_logger.log_operation(
                operation="ai_clean_file",
                success=False,
                details={"filepath": filepath, "duration_ms": round(duration_ms, 2)},
                error=str(e),
                command_text=command_text
            )
            return {"success": False, "message": f"AI清洗失败: {str(e)}"}